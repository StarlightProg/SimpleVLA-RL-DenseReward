from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "dense_reward_validation_report.docx"
FIGURES = [
    Path("/Users/daeron/Downloads/photo_2026-08-07 17.18.25.jpeg"),
    Path("/Users/daeron/Downloads/photo_2026-08-07 17.18.27.jpeg"),
    Path("/Users/daeron/Downloads/photo_2026-08-07 17.18.29.jpeg"),
]

BLUE = "1F4E79"
TEAL = "0F766E"
GREEN = "2E7D32"
ORANGE = "C2410C"
PURPLE = "6D46D7"
GRAY = "5B6573"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
LIGHT_YELLOW = "FFF6D8"


def set_cell_shading(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.tcPr.tcW
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def style_table(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for i, cell in enumerate(table.rows[0].cells):
        set_cell_shading(cell, LIGHT_BLUE)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(BLUE)
            run.font.size = Pt(9)
    set_repeat_table_header(table.rows[0])
    for row in table.rows[1:]:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(9)


def add_text(doc, text, bold=False, color=None, size=11, after=6, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(10.5)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.bold = True
    r.font.size = Pt(16 if level == 1 else 12.5)
    r.font.color.rgb = RGBColor.from_string(BLUE if level == 1 else TEAL)
    return p


def add_callout(doc, label, text, fill=LIGHT_YELLOW):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(label + " ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)
    r.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(doc, path, caption):
    doc.add_picture(str(path), width=Inches(6.35))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(GRAY)


def setup_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)
    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    styles["Normal"].font.size = Pt(11)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("SimpleVLA-RL DenseReward | Validation report")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Internal experiment report | 7 August 2026")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(GRAY)


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Dense Reward Validation Report")
    r.font.name = "Calibri"
    r.font.size = Pt(25)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(BLUE)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(16)
    r2 = p2.add_run("LIBERO-Spatial: comparison with the original terminal-only reward")
    r2.font.size = Pt(13)
    r2.font.color.rgb = RGBColor.from_string(GRAY)
    metadata = doc.add_table(rows=3, cols=2)
    metadata.cell(0, 0).text = "Project revision"
    metadata.cell(0, 1).text = "SimpleVLA-RL-DenseReward @ 86111de (validation optimization)"
    metadata.cell(1, 0).text = "Evidence"
    metadata.cell(1, 1).text = "16 validation logs; fixed 300-rollout evaluation per checkpoint; three supplied plots"
    metadata.cell(2, 0).text = "Starting checkpoint"
    metadata.cell(2, 1).text = "≈50.0% validation success before RL (per run-owner note and figures)"
    set_table_geometry(metadata, [2100, 7260])
    for row in metadata.rows:
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(BLUE)
            run.font.size = Pt(9.5)
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(9.5)


def add_results_table(doc):
    add_heading(doc, "Results", 1)
    add_text(doc, "All values are validation success rates on the same fixed LIBERO-Spatial validation set. Each checkpoint evaluation contains 300 rollouts (10 tasks × 30 trials).", color=GRAY, size=9.5, after=7)
    rows = [
        ["Terminal only (coef=5)", "61.7%", "step 170", "+11.7 pp", "Reference baseline"],
        ["Dense 0.05/0.05 + terminal 5", "65.3%", "step 160", "+15.3 pp", "+3.6 pp vs terminal"],
        ["Dense clipped + terminal 5", "64.7%", "step 50", "+14.7 pp", "Filtered trajectory set"],
        ["Dense 0.05/0.05 + terminal 1", "61.0%", "step 110", "+11.0 pp", "Different terminal scale"],
        ["Phase-only + terminal 1", "63.7%", "step 90", "+13.7 pp", "Peak; clipped at 0.05"],
    ]
    table = doc.add_table(rows=1, cols=5)
    headers = ["Reward variant", "Best success", "Checkpoint", "Δ from 50.0%", "Interpretation"]
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value
    style_table(table, [3050, 1200, 1200, 1350, 2560])


def add_trajectory_table(doc):
    add_heading(doc, "Terminal-only baseline vs. dense reward with terminal coefficient 5", 2)
    rows = [
        ["40", "61.0%", "61.8%", "+0.8 pp"],
        ["80", "60.7%", "63.0%", "+2.3 pp"],
        ["110", "61.3%", "63.7%", "+2.4 pp"],
        ["160/170", "61.7% (170)", "65.3% (160)", "+3.6 pp"],
    ]
    table = doc.add_table(rows=1, cols=4)
    for cell, value in zip(table.rows[0].cells, ["Approx. checkpoint", "Terminal only", "Dense 0.05/0.05 + terminal 5", "Dense − terminal"]):
        cell.text = value
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value
    style_table(table, [1500, 1850, 3600, 2410])


def main():
    doc = Document()
    setup_doc(doc)
    add_title(doc)
    add_heading(doc, "Executive summary", 1)
    add_text(doc, "The dense reward with terminal coefficient 5 is the strongest tested reward under this validation protocol. It reached 65.3% success at checkpoint 160, exceeding the terminal-only reference’s best observed 61.7% at checkpoint 170 by 3.6 percentage points (pp). From the reported 50.0% pre-RL checkpoint, this is a +15.3 pp gain versus +11.7 pp for terminal-only.")
    add_text(doc, "At the three directly aligned checkpoints (40, 80, and 110), dense reward is higher than terminal-only by 0.8, 2.3, and 2.4 pp. The observed advantage therefore appears consistently in the sampled checkpoints and grows over the measured training range.")
    add_callout(doc, "Decision: ", "Use dense 0.05/0.05 + terminal 5 as the next candidate configuration. Keep terminal success at coefficient 5 for the primary comparison, then confirm the result with multiple training seeds and a held-out evaluation seed set.", LIGHT_BLUE)
    add_heading(doc, "Scope and protocol", 1)
    add_bullet(doc, "Benchmark: LIBERO-Spatial; success rate is the mean of 300 fixed validation rollouts per evaluated checkpoint.")
    add_bullet(doc, "Model and validation settings in the logs: OpenVLA-OFT, LoRA rank 16 on llm-projector; validation uses libero_spatial_no_noops.")
    add_bullet(doc, "This report compares reward variants within the supplied runs. It does not claim reproduction of the published SimpleVLA-RL results because the starting checkpoint and training recipe differ.")
    add_results_table(doc)
    add_trajectory_table(doc)
    add_heading(doc, "Interpretation by reward", 1)
    add_heading(doc, "1. Original terminal-only reward (coefficient 5)", 2)
    add_text(doc, "This is the reference condition: reward is driven only by episode completion, with terminal coefficient 5. It improves from the reported 50.0% starting point to 60.7–61.7%, but the four observed checkpoints form a near-flat plateau after step 40. Its best gain is +11.7 pp.")
    add_heading(doc, "2. Dense 0.05/0.05 + terminal 5", 2)
    add_text(doc, "This is the main dense-reward condition described by the run owner: 0.05 for within-phase progress and 0.05 for completing a phase, added to terminal reward 5. It rises from 61.8% at step 40 to 65.3% at step 160 without a sampled regression. Its best gain from the starting checkpoint is 15.3 pp, 30.8% larger than the terminal-only gain (15.3 / 11.7 − 1).")
    add_heading(doc, "3. Dense clipped + terminal 5", 2)
    add_text(doc, "This variant reaches 64.7% at step 50, close to the best 65.3% dense run. It trains on a selected trajectory set that excludes fully successful and fully failed trajectories, and the run owner reports approximately four days of training despite only 50 saved steps. Therefore, its 50-step value is evidence of a promising result, but it is not an apples-to-apples claim of faster learning by step count or wall-clock time.")
    add_heading(doc, "4. Terminal coefficient 1 ablations", 2)
    add_text(doc, "Dense 0.05/0.05 + terminal 1 peaks at 61.0%, while the phase-only condition peaks at 63.7% before falling to 62.3% at step 110. These are useful exploratory results but are confounded by the lower terminal coefficient. The run owner notes that phase-only shaping remained clipped at 0.05 because the clip was not increased; this limits conclusions about the intended phase-reward magnitude.")
    add_heading(doc, "What the data support—and what they do not", 1)
    add_bullet(doc, "Supported: the dense 0.05/0.05 + terminal 5 series is above terminal-only at every directly aligned measured checkpoint and has the best final observed success rate.")
    add_bullet(doc, "Supported: dense clipping produces a high success rate under its own filtered-trajectory procedure, but with a different compute/selection profile.")
    add_bullet(doc, "Not yet supported: a statistically confirmed improvement. With 300 binary rollouts, individual 95% half-widths in the logs are approximately ±5.4–5.6 pp; the 3.6 pp best-checkpoint gap requires repeated training seeds and paired or independent statistical testing.")
    add_bullet(doc, "Not yet supported: a causal conclusion that lowering terminal reward from 5 to 1 causes instability. The observed pattern is consistent with that hypothesis, but reward scale and dense-reward composition changed simultaneously.")
    add_heading(doc, "Recommended next experiment", 1)
    numbered = [
        "Run terminal-only (coefficient 5) and dense 0.05/0.05 + terminal 5 with at least three matched training seeds from the same ≈50% starting checkpoint.",
        "Evaluate every selected checkpoint on the same 300-rollout fixed set and an additional held-out seed set; retain per-episode success IDs to permit paired comparisons.",
        "Match the dense-clipped and unfiltered conditions by wall-clock budget and record the number of collected, retained, fully successful, and fully failed trajectories.",
        "Run a controlled terminal-scale ablation: terminal 1 versus 5 while keeping dense weights, clip, batch composition, and all other settings fixed.",
        "For phase-only shaping, explicitly set and record the intended clip; then repeat its best configuration with matched seeds.",
    ]
    for item in numbered:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        r.font.size = Pt(10.5)
    doc.add_page_break()
    add_heading(doc, "Appendix A — supplied validation plots", 1)
    captions = [
        "Figure A1. Terminal-only versus the two terminal=5 dense-reward variants.",
        "Figure A2. Validation success across all supplied ablations.",
        "Figure A3. Best evaluated checkpoint for each ablation.",
    ]
    for index, (figure, caption) in enumerate(zip(FIGURES, captions)):
        if index:
            doc.add_page_break()
        add_figure(doc, figure, caption)
    doc.add_page_break()
    add_heading(doc, "Appendix B — source inventory", 1)
    add_text(doc, "Primary quantitative source: /Users/daeron/Downloads/vaidation_runs/plots/validation_metrics_summary.csv and the 16 corresponding validation logs in the same directory. The log evidence reports the final per-checkpoint aggregate validation metric and rollout count.")
    add_text(doc, "Qualitative configuration and runtime notes: supplied by the run owner in the task request (starting success ≈0.5, dense/phase definitions, and dense-clipped runtime). These notes are labelled as run-owner descriptions where they are not independently reconstructible from the validation logs.")
    add_text(doc, "Codebase inspected: SimpleVLA-RL-DenseReward at git revision 86111de (2026-07-14). Relevant implementation locations include verl/utils/subgoal_reward/dense_reward.py, verl/utils/subgoal_reward/engine.py, and verl/trainer/ppo/ray_trainer.py.")
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
